import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
import os
import json
from docx import Document
from database import Database
from ocr_processor import OcrProcessor

class Aplikace:
    def __init__(self, root):
        self.db = Database()
        self.root = root
        self.root.title("Registr chovu králíků")
        self.root.geometry("1000x750")

        # --- HORNÍ PANEL ---
        top_frame = tk.Frame(root, pady=10)
        top_frame.pack(fill="x", padx=10)

        tk.Button(top_frame, text="Přidat králíka", command=self.novy_kralik_okno, bg="#4CAF50", fg="white").pack(side="left", padx=5)
        tk.Button(top_frame, text="Import z AI (Text)", command=self.otevri_import_okno, bg="#2196F3", fg="white").pack(side="left", padx=5)

        tk.Button(top_frame, text="Nápověda", command=self.ukaz_napovedu_ai, 
                  bg="#607D8B", fg="white", width=10).pack(side="left", padx=5)
        
        tk.Label(top_frame, text="Hledat:").pack(side="left", padx=(20, 5))
        self.search_var = tk.StringVar()
        self.search_var.trace("w", lambda *args: self.aktualizuj_tabulku(self.search_var.get()))
        tk.Entry(top_frame, textvariable=self.search_var, width=30).pack(side="left")

        # --- TABULKA ---
        self.tree = ttk.Treeview(root, columns=("id_sql", "tetovani", "pohlavi", "plemeno"), show="headings")
        self.tree.heading("id_sql", text="ID")
        self.tree.heading("tetovani", text="Tetování (L/P)")
        self.tree.heading("pohlavi", text="Pohlaví")
        self.tree.heading("plemeno", text="Plemeno")
        
        self.tree.column("id_sql", width=0, stretch=False)
        self.tree.column("tetovani", width=150)
        self.tree.pack(fill="both", expand=True, padx=10, pady=10)

        self.tree.bind("<Double-1>", self.otevri_detail)
        tk.Button(root, text="Smazat vybraného", command=self.smazat_z_tabulky, bg="#f44336", fg="white").pack(pady=5)

        self.aktualizuj_tabulku()

    def aktualizuj_tabulku(self, filtr=""):
        for i in self.tree.get_children(): self.tree.delete(i)
        seznam = self.db.ziskej_vsechny_kraliky()
        filtr = filtr.lower()
        for k in seznam:
            tetovani = f"{k['levo']} / {k['pravo']}"
            if not filtr or filtr in k['plemeno'].lower() or filtr in tetovani.lower():
                self.tree.insert("", "end", values=(k['id'], tetovani, k['pohlavi'], k['plemeno']))

    def smazat_z_tabulky(self):
        vyber = self.tree.selection()
        if not vyber: return
        v_data = self.tree.item(vyber[0])['values']
        if messagebox.askyesno("Smazat", f"Opravdu smazat králíka {v_data[1]}?"):
            self.db.smaz_kralika(v_data[0])
            self.aktualizuj_tabulku()

    def otevri_import_okno(self):
        import_okno = tk.Toplevel(self.root)
        import_okno.title("Vložit data z AI")
        import_okno.geometry("600x500")
        tk.Label(import_okno, text="Vložte JSON kód od AI:", font=("Arial", 10, "bold")).pack(pady=10)
        text_area = tk.Text(import_okno, height=20, width=70)
        text_area.pack(padx=10, pady=10)

        def potvrdit():
            obsah = text_area.get("1.0", tk.END).strip()
            try:
                data = json.loads(obsah)
                # Posíláme celý JSON, ne jen data.get('hlavni_kralik')
                # Tím umožníme uložení předků
                self.db.uloz_kralika(data.get('hlavni_kralik', data)) 
                
                self.aktualizuj_tabulku()
                import_okno.destroy()
                messagebox.showinfo("Hotovo", "Králík i jeho předci byli úspěšně nahráni.")
            except Exception as e:
                messagebox.showerror("Chyba", f"Neplatný formát dat: {e}")

        tk.Button(import_okno, text="Uložit", command=potvrdit, bg="#4CAF50", fg="white").pack()

    def novy_kralik_okno(self):
        self.otevri_editor(None)

    def otevri_detail(self, event):
        vyber = self.tree.selection()
        if not vyber: return
        k_id_sql = self.tree.item(vyber[0])['values'][0]
        kralik_data = self.db.ziskej_kralika_podle_id(k_id_sql)
        self.otevri_editor(kralik_data)

    def otevri_editor(self, kralik):
        editor = tk.Toplevel(self.root)
        editor.title("Karta králíka")
        editor.geometry("700x850")

        notebook = ttk.Notebook(editor)
        notebook.pack(fill="both", expand=True, padx=10, pady=10)

        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="Základní údaje")

        # Seznam všech polí
        fields = [
            ("Levé ucho (L. u.):", "levo"), ("Pravé ucho (P. u.):", "pravo"),
            ("Jméno:", "jmeno"), ("Pohlaví (1,0/0,1):", "pohlavi"),
            ("Plemeno:", "plemeno"), ("Barva:", "barva"),
            ("Datum vrhu:", "datum_narozeni"), ("Hmotnost (kg):", "hmotnost"),
            ("Chovatel:", "chovatel"), ("Adresa chovatele:", "adresa"),
            ("Narozených (ks):", "naroz_ks"), ("Odchovaných (ks):", "odchov_ks"),
            ("Registrovaných (ks):", "registr_ks")
        ]
        
        entries = {}
        for i, (label_text, attr) in enumerate(fields):
            tk.Label(tab1, text=label_text).grid(row=i, column=0, sticky="e", padx=5, pady=3)
            ent = tk.Entry(tab1, width=35)
            ent.grid(row=i, column=1, sticky="w", padx=5, pady=3)
            
            # Vložení dat, pokud existují
            if kralik:
                hodnota = kralik.get(attr)
                ent.insert(0, str(hodnota) if hodnota is not None else "")
            
            # Funke Enteru pro přeskok
            ent.bind("<Return>", lambda e: e.widget.tk_focusNext().focus())
            entries[attr] = ent

        tk.Button(tab1, text="Sken (Tesseract)", command=lambda: self.nacti_z_fotky(entries), bg="#9C27B0", fg="white").grid(row=0, column=2, padx=10)

        # --- ZÁLOŽKA 2: PŘEDCI ---
        tab2 = ttk.Frame(notebook)
        notebook.add(tab2, text="Předci")

        # 1. Tlačítko pro skenování 
        tk.Button(tab2, text="Skenovat zadní stranu (14 předků)", 
                  command=lambda: self.skenuj_predky_lokalne(kralik, entries),
                  bg="#FF9800", fg="white", font=("Arial", 9, "bold")).pack(pady=15)

        # Pomocná funkce pro zobrazení tetování místo suchého čísla ID
        def dej_tetovani(k_id):
            if not k_id: return "Nezadáno"
            k = self.db.ziskej_kralika_podle_id(k_id)
            return f"{k['levo']} / {k['pravo']}" if k else "Nezadáno"

        # Dočasné proměnné pro uložení ID vybraných rodičů
        self.temp_otec_id = kralik.get('otec_id') if kralik else None
        self.temp_matka_id = kralik.get('matka_id') if kralik else None

        # --- SEKCE OTEC ---
        frame_o = tk.LabelFrame(tab2, text="Otec (1,0)", padx=10, pady=10)
        frame_o.pack(fill="x", padx=15, pady=5)
        
        lbl_otec_text = tk.Label(frame_o, text=dej_tetovani(self.temp_otec_id), 
                                 font=("Arial", 10, "bold"), fg="#2196F3")
        lbl_otec_text.pack(side="left", padx=10)

        def zmenit_otce():
            v = self.vyber_kralika_dialog("1,0") # Dialog ukáže jen samce
            if v["id"]:
                self.temp_otec_id = v["id"]
                lbl_otec_text.config(text=v["tetovani"])

        tk.Button(frame_o, text="Vybrat ze seznamu", command=zmenit_otce).pack(side="right")

        # --- SEKCE MATKA ---
        frame_m = tk.LabelFrame(tab2, text="Matka (0,1)", padx=10, pady=10)
        frame_m.pack(fill="x", padx=15, pady=5)

        lbl_matka_text = tk.Label(frame_m, text=dej_tetovani(self.temp_matka_id), 
                                  font=("Arial", 10, "bold"), fg="#E91E63")
        lbl_matka_text.pack(side="left", padx=10)

        def zmenit_matku():
            v = self.vyber_kralika_dialog("0,1") # Dialog ukáže jen samice
            if v["id"]:
                self.temp_matka_id = v["id"]
                lbl_matka_text.config(text=v["tetovani"])

        tk.Button(frame_m, text="Vybrat ze seznamu", command=zmenit_matku).pack(side="right")

        def ulozit():
            # 1. Posbíráme textová data z políček (uši, plemeno, chovatel atd.)
            data = {attr: entries[attr].get() for _, attr in fields}
            
            # 2. Uložíme králíka do SQL databáze a získáme jeho ID
            k_id = self.db.uloz_kralika(data)
            
            if k_id:
                # 3. Použijeme ID rodičů, která jsou uložená v temp proměnných
                self.db.nastav_rodice(k_id, self.temp_otec_id, self.temp_matka_id)
                
                # 4. Obnovíme zobrazení a zavřeme okno
                self.aktualizuj_tabulku()
                editor.destroy()
                messagebox.showinfo("Hotovo", "Králík byl úspěšně uložen do databáze.")
            else:
                messagebox.showerror("Chyba", "Nepodařilo se uložit data do SQL.")

        # --- TLAČÍTKA NA SPODU OKNA ---
        tk.Button(editor, text="Uložit králíka", command=ulozit, bg="#4CAF50", fg="white", pady=10).pack(pady=5)

        # Tisková sekce (Jen pro existující králíky)
        if kralik:
            print_frame = tk.LabelFrame(editor, text="Tisk dokumentů", padx=10, pady=10)
            print_frame.pack(fill="x", padx=10, pady=10)
            
            tk.Button(print_frame, text="Generovat RODOKMEN", 
                      command=lambda: self.generovat_jen_rodokmen(kralik), 
                      bg="#2196F3", fg="white", width=40).pack(pady=2)
            
            tk.Button(print_frame, text="Generovat POTVRZENÍ", 
                      command=lambda: self.generovat_jen_potvrzeni(kralik), 
                      bg="#FF9800", fg="white", width=40).pack(pady=2)
            
            tk.Button(print_frame, text="KOMPLETNÍ DOKUMENT", 
                      command=lambda: self.generovat_kompletni_dokument(kralik), 
                      bg="#4CAF50", fg="white", font=("Arial", 9, "bold"), width=40, pady=5).pack(pady=10)       

    # --- LOGIKA TISKU ---
    def _ziskej_nahrady_rodokmenu(self, kralik):
        """Pomocná funkce pro vyhledání předků v SQL a přípravu značek {{...}}"""
        def get_anc_data(k_id, prefix, is_parent=False):
            # Pokud ID chybí nebo králík v DB není, vrátíme prázdné značky
            if not k_id: 
                res = {f"{{{{{prefix}_l}}}}": "", f"{{{{{prefix}_p}}}}": "", f"{{{{{prefix}_ch}}}}": ""}
                if is_parent:
                    res.update({f"{{{{{prefix}_pl}}}}": "", f"{{{{{prefix}_dv}}}}": ""})
                return res

            k = self.db.ziskej_kralika_podle_id(k_id)
            if not k: 
                res = {f"{{{{{prefix}_l}}}}": "", f"{{{{{prefix}_p}}}}": "", f"{{{{{prefix}_ch}}}}": ""}
                if is_parent:
                    res.update({f"{{{{{prefix}_pl}}}}": "", f"{{{{{prefix}_dv}}}}": ""})
                return res
            
            # Základní údaje pro všechny předky
            res = {
                f"{{{{{prefix}_l}}}}": k['levo'] or "",
                f"{{{{{prefix}_p}}}}": k['pravo'] or "",
                f"{{{{{prefix}_ch}}}}": f"{k['chovatel']}, {k['adresa']}" if k['chovatel'] else ""
            }
            
            # Údaje navíc pouze pro přímé rodiče (otec_pl, otec_dv, atd.)
            if is_parent:
                res[f"{{{{{prefix}_pl}}}}"] = k['plemeno'] or ""
                res[f"{{{{{prefix}_dv}}}}"] = k['datum_narozeni'] or ""

            # ID rodičů pro další generaci
            res['o'], res['m'] = k['otec_id'], k['matka_id']
            return res

        # Údaje o hlavním králíkovi
        nahrady = {
            "{{jmeno}}": kralik['jmeno'] or "", 
            "{{levo}}": kralik['levo'] or "", 
            "{{pravo}}": kralik['pravo'] or "",
            "{{plemeno}}": kralik['plemeno'] or "", 
            "{{datum_vrhu}}": kralik['datum_narozeni'] or "",
            "{{chovatel}}": kralik['chovatel'] or "", 
            "{{adresa}}": kralik['adresa'] or "",
            "{{hmotnost}}": str(kralik['hmotnost'] if kralik['hmotnost'] else ""),
            "{{v_nar}}": str(kralik['naroz_ks'] if kralik['naroz_ks'] is not None else "0"),
            "{{v_odch}}": str(kralik['odchov_ks'] if kralik['odchov_ks'] is not None else "0"),
            "{{v_reg}}": str(kralik['registr_ks'] if kralik['registr_ks'] is not None else "0")
        }

        # 1. Generace (Rodiče) - Zde nastavujeme is_parent=True
        o = get_anc_data(kralik['otec_id'], "otec", is_parent=True)
        m = get_anc_data(kralik['matka_id'], "matka", is_parent=True)
        
        # 2. Generace (Prarodiče)
        oo = get_anc_data(o.get('o'), "oo")
        mo = get_anc_data(o.get('m'), "mo")
        om = get_anc_data(m.get('o'), "om")
        mm = get_anc_data(m.get('m'), "mm")

        # 3. Generace (Praprarodiče)
        ooo = get_anc_data(oo.get('o'), "ooo")
        moo = get_anc_data(oo.get('m'), "moo")
        omo = get_anc_data(mo.get('o'), "omo")
        mmo = get_anc_data(mo.get('m'), "mmo")
        oom = get_anc_data(om.get('o'), "oom")
        mom = get_anc_data(om.get('m'), "mom")
        omm = get_anc_data(mm.get('o'), "omm")
        mmm = get_anc_data(mm.get('m'), "mmm")

        # Spojíme všechny nasbírané značky do výsledného slovníku
        vsechny_generace = [o, m, oo, mo, om, mm, ooo, moo, omo, mmo, oom, mom, omm, mmm]
        for gen in vsechny_generace:
            for k, v in gen.items():
                if k.startswith("{{"):
                    nahrady[k] = v

        return nahrady

    def generovat_jen_rodokmen(self, kralik):
        try:
            doc = Document("rodokmen.docx")
            nahrady = self._ziskej_nahrady_rodokmenu(kralik)
            self._proved_nahrazeni(doc, nahrady)
            nazev = f"Rodokmen_{kralik['levo']}_{kralik['pravo']}.docx".replace("/", "-")
            doc.save(nazev)
            os.startfile(nazev)
        except Exception as e: messagebox.showerror("Chyba", f"Rodokmen: {e}")

    def generovat_jen_potvrzeni(self, kralik):
        p_id = simpledialog.askstring("Partner", "Zadejte SQL ID partnera (číslo):")
        if not p_id: return
        partner = self.db.ziskej_kralika_podle_id(int(p_id))
        
        try:
            doc = Document("rodokmen.docx")
            nahrady = self._priprav_data_potvrzeni(kralik, partner)
            self._proved_nahrazeni(doc, nahrady)
            nazev = f"Potvrzeni_{kralik['levo']}.docx".replace("/", "-")
            doc.save(nazev)
            os.startfile(nazev)
        except Exception as e: messagebox.showerror("Chyba", f"Potvrzení: {e}")

    def generovat_kompletni_dokument(self, kralik):
        p_vstup = simpledialog.askstring("Partner", "Zadejte ID (číslo) nebo tetování partnera (např. C1-2 / 10-20):")
        if not p_vstup: return
        
        # Inteligentní vyhledání partnera
        if p_vstup.isdigit():
            partner = self.db.ziskej_kralika_podle_id(int(p_vstup))
        else:
            partner = self.db.najdi_podle_tetovani(p_vstup)

        if not partner:
            messagebox.showwarning("Chyba", "Partner nebyl v databázi nalezen.")
            return

        datum_kryti = simpledialog.askstring("Krytí", "Datum krytí:", initialvalue="1. 1. 2026")

        try:
            doc = Document("rodokmen.docx")
            nahrady = self._ziskej_nahrady_rodokmenu(kralik)
            potvrzeni = self._priprav_data_potvrzeni(kralik, partner)
            nahrady.update(potvrzeni)
            nahrady["{{k_datum}}"] = datum_kryti
            
            self._proved_nahrazeni(doc, nahrady)
            nazev = f"Kompletni_{kralik['levo']}.docx".replace("/", "-").replace(" ", "")
            doc.save(nazev)
            os.startfile(nazev)
        except Exception as e: 
            messagebox.showerror("Chyba", f"Generování selhalo: {e}")

    def _priprav_data_potvrzeni(self, kralik, partner):
        
        s_hlavni = kralik if kralik['pohlavi'] == "1,0" else partner
        m_hlavni = kralik if kralik['pohlavi'] == "0,1" else partner
        
        return {
            "{{m_majitel}}": s_hlavni['chovatel'] if s_hlavni else "",
            "{{m_adresa}}": s_hlavni['adresa'] if s_hlavni else "",
            "{{s_majitel}}": m_hlavni['chovatel'] if m_hlavni else "",
            "{{s_adresa}}": m_hlavni['adresa'] if m_hlavni else "",
            "{{p_plemeno}}": m_hlavni['plemeno'] if m_hlavni else "",
            "{{p_barva}}": m_hlavni['barva'] if m_hlavni else "",
            "{{v_nar}}": str(m_hlavni['naroz_ks'] if m_hlavni and m_hlavni['naroz_ks'] is not None else "0"),
            "{{v_odch}}": str(m_hlavni['odchov_ks'] if m_hlavni and m_hlavni['odchov_ks'] is not None else "0"),
            "{{v_reg}}": str(m_hlavni['registr_ks'] if m_hlavni and m_hlavni['registr_ks'] is not None else "0")
        }

    def _proved_nahrazeni(self, doc, nahrady):
        for p in doc.paragraphs:
            for k, v in nahrady.items():
                if k in p.text: p.text = p.text.replace(k, str(v) if v else "")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for k, v in nahrady.items():
                        if k in cell.text: cell.text = cell.text.replace(k, str(v) if v else "")

    def nacti_z_fotky(self, entries):
        cesta = filedialog.askopenfilename(filetypes=[("Obrázky", "*.png *.jpg *.jpeg")])
        if not cesta: return
        try:
            if not hasattr(self, 'ocr'): self.ocr = OcrProcessor()
            data = self.ocr.nacti_data_z_obrazku(cesta)
            for k, v in data.items():
                if k in entries:
                    entries[k].delete(0, tk.END)
                    entries[k].insert(0, v)
        except Exception as e: messagebox.showerror("OCR", str(e))

    def skenuj_predky_lokalne(self, kralik, entries):
        cesta = filedialog.askopenfilename(filetypes=[("Obrázky", "*.png *.jpg *.jpeg")])
        if not cesta: return
        
        try:
            if not hasattr(self, 'ocr'): self.ocr = OcrProcessor()
            tattoos = self.ocr.nacti_zadni_stranu(cesta)
            
            if len(tattoos) >= 2:
                # Otec a Matka jsou první dvě nalezená tetování
                messagebox.showinfo("Sken předků", 
                    f"Nalezena tetování:\nOtec: {tattoos[0]}\nMatka: {tattoos[1]}\n\n"
                    "Byl vytvořen základ pro rodokmen.")
                
            else:
                messagebox.showwarning("Sken", "Nepodařilo se najít dostatek tetování.")
        except Exception as e:
            messagebox.showerror("Chyba skenu", str(e))

    def skenuj_zadni_stranu_akce(self, kralik_id):
        cesta = filedialog.askopenfilename(title="Vyberte zadní stranu rodokmenu")
        if not cesta: return
        
        if not hasattr(self, 'ocr'): self.ocr = OcrProcessor()
        seznam = self.ocr.nacti_zadni_stranu(cesta)
        
        if len(seznam) >= 4:
            if self.db.uloz_14_predku(kralik_id, seznam):
                messagebox.showinfo("Úspěch", f"Načteno {len(seznam)} tetování a propojen Otec a Matka.")
                self.aktualizuj_tabulku()
        else:
            messagebox.showwarning("Sken", f"Nalezeno pouze {len(seznam)} tetování. Zkuste lepší fotku.")

    def vyber_kralika_dialog(self, pohlavi_hledane=None):
        """Otevře okno se seznamem králíků a vrátí ID vybraného jedince."""
        vysledek = {"id": None, "tetovani": "Nezadáno"}
        
        okno = tk.Toplevel(self.root)
        okno.title("Vyberte králíka")
        okno.geometry("500x400")
        okno.grab_set()

        tree = ttk.Treeview(okno, columns=("id", "tetovani", "plemeno"), show="headings")
        tree.heading("id", text="ID"); tree.heading("tetovani", text="Tetování"); tree.heading("plemeno", text="Plemeno")
        tree.column("id", width=50); tree.pack(fill="both", expand=True, padx=10, pady=10)

        # Načtení králíků z SQL
        vsechny = self.db.ziskej_vsechny_kraliky()
        for k in vsechny:
            if pohlavi_hledane and k['pohlavi'] != pohlavi_hledane:
                continue
            tree.insert("", "end", values=(k['id'], f"{k['levo']} / {k['pravo']}", k['plemeno']))

        def potvrdit(event=None):
            vyber = tree.selection()
            if vyber:
                v_data = tree.item(vyber[0])['values']
                vysledek["id"] = v_data[0]
                vysledek["tetovani"] = v_data[1]
                okno.destroy()

        tree.bind("<Double-1>", potvrdit)
        tk.Button(okno, text="Vybrat", command=potvrdit, bg="#4CAF50", fg="white").pack(pady=10)
        
        self.root.wait_window(okno)
        return vysledek
    
    def ukaz_napovedu_ai(self):
        """Otevře okno s promptem pro AI a umožní jeho zkopírování."""
        napoveda_okno = tk.Toplevel(self.root)
        napoveda_okno.title("Nápověda k AI importu")
        napoveda_okno.geometry("500x400")
        
        instrukce = (
            "Pro nejlepší výsledek pošli AI chatbotovi fotku rodokmenu\n"
            "a k ní přilož tento příkaz (prompt):\n"
        )
        tk.Label(napoveda_okno, text=instrukce, font=("Arial", 10)).pack(pady=10)

        # Textové pole s promptem
        prompt_text = (
            "Analyzuj tento rodokmen. Vytvoř JSON pro můj program. "
            "Musíš u všech králíků (hlavní i 14 předků) uvést: "
            "'levo', 'pravo', 'pohlavi' (1,0 nebo 0,1), 'chovatel' a 'adresa'. "
            "U hlavního králíka a jeho rodičů přidej i 'plemeno' a 'datum_narozeni'. "
            "Pokud údaj vidíš, musí být v JSONu. Odpověz pouze čistým JSON kódem."
        )
        
        text_box = tk.Text(napoveda_okno, height=8, padx=10, pady=10, wrap="word")
        text_box.insert("1.0", prompt_text)
        text_box.config(state="disabled") 
        text_box.pack(padx=20, pady=5)

        def kopirovat():
            self.root.clipboard_clear()
            self.root.clipboard_append(prompt_text)
            messagebox.showinfo("Zkopírováno", "Prompt byl zkopírován do schránky.")

        tk.Button(napoveda_okno, text="Zkopírovat prompt", command=kopirovat, 
                  bg="#2196F3", fg="white", pady=5).pack(pady=10)
        
        tk.Label(napoveda_okno, text="Poté vygenerovaný kód vlož přes tlačítko 'Import z AI'.", 
                 font=("Arial", 8, "italic")).pack(pady=5)

if __name__ == "__main__":
    root = tk.Tk()
    app = Aplikace(root)
    root.mainloop()